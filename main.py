import os
import re
import json
import uuid
import logging
from io import BytesIO
from flask import Flask, request, render_template, send_from_directory, url_for, flash, redirect, session, jsonify, send_file
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from weasyprint import HTML, CSS

# Configuração de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('aibnt_app')

# Configurações do aplicativo
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'uploads')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'gdoc'}
MAX_CONTENT_LENGTH = 20 * 1024 * 1024  # 20MB limite

app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.secret_key = os.urandom(24)  # Necessário para mensagens flash e sessões

# Garantir que a pasta de uploads exista
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Simulação de banco de dados de usuários (em produção, usar um banco de dados real)
users_db = {}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Classe de formatação ABNT completa
class ABNTFormatter:
    """
    Classe responsável pela formatação de documentos conforme as normas ABNT 2023.
    """
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.pdf', '.txt']
        
    def format_document(self, input_path, output_dir=None):
        """
        Formata um documento de acordo com as normas ABNT 2023.
        
        Args:
            input_path (str): Caminho do arquivo de entrada
            output_dir (str, optional): Diretório de saída. Se None, usa o mesmo diretório do arquivo de entrada.
            
        Returns:
            str: Caminho do arquivo formatado
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Arquivo não encontrado: {input_path}")
        
        # Determinar extensão do arquivo
        _, ext = os.path.splitext(input_path)
        ext = ext.lower()
        
        if ext not in self.supported_extensions:
            raise ValueError(f"Formato de arquivo não suportado: {ext}. Formatos suportados: {', '.join(self.supported_extensions)}")
        
        # Definir diretório de saída
        if output_dir is None:
            output_dir = os.path.dirname(input_path)
        
        # Criar diretório de saída se não existir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Definir caminho de saída
        filename = os.path.basename(input_path)
        name, _ = os.path.splitext(filename)
        
        # Formatar documento de acordo com a extensão
        try:
            if ext == '.docx':
                output_path = os.path.join(output_dir, f"{name}_ABNT.docx")
                return self._format_docx(input_path, output_path)
            elif ext == '.pdf':
                output_path = os.path.join(output_dir, f"{name}_ABNT.pdf")
                return self._format_pdf(input_path, output_path)
            elif ext == '.txt':
                output_path = os.path.join(output_dir, f"{name}_ABNT.docx")
                return self._format_txt_to_docx(input_path, output_path)
            else:
                raise ValueError(f"Formato não suportado: {ext}")
        except Exception as e:
            logger.error(f"Erro ao formatar documento: {str(e)}")
            raise
    
    def _format_docx(self, input_path, output_path):
        """
        Formata um documento DOCX de acordo com as normas ABNT 2023.
        
        Args:
            input_path (str): Caminho do arquivo de entrada
            output_path (str): Caminho do arquivo de saída
            
        Returns:
            str: Caminho do arquivo formatado
        """
        try:
            # Abrir documento
            doc = Document(input_path)
            
            # Configurar estilos ABNT
            self._apply_abnt_styles(doc)
            
            # Aplicar formatação ABNT ao conteúdo
            self._apply_abnt_formatting_to_docx(doc)
            
            # Salvar documento formatado
            doc.save(output_path)
            
            logger.info(f"Documento DOCX formatado com sucesso: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Erro ao formatar documento DOCX: {str(e)}")
            raise
    
    def _format_pdf(self, input_path, output_path):
        """
        Formata um documento PDF de acordo com as normas ABNT 2023.
        
        Args:
            input_path (str): Caminho do arquivo de entrada
            output_path (str): Caminho do arquivo de saída
            
        Returns:
            str: Caminho do arquivo formatado
        """
        try:
            # Extrair texto do PDF
            reader = PdfReader(input_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            
            # Criar documento DOCX temporário
            temp_docx = os.path.join(os.path.dirname(output_path), f"temp_{uuid.uuid4()}.docx")
            doc = Document()
            
            # Configurar estilos ABNT
            self._apply_abnt_styles(doc)
            
            # Adicionar texto extraído
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style = 'Normal'
            
            # Aplicar formatação ABNT ao conteúdo
            self._apply_abnt_formatting_to_docx(doc)
            
            # Salvar documento DOCX temporário
            doc.save(temp_docx)
            
            # Converter DOCX para PDF
            docx_output_path = output_path.replace('.pdf', '.docx')
            os.rename(temp_docx, docx_output_path)
            
            # Também retornar o DOCX para download
            logger.info(f"Documento PDF convertido para DOCX com formatação ABNT: {docx_output_path}")
            return docx_output_path
        except Exception as e:
            logger.error(f"Erro ao formatar documento PDF: {str(e)}")
            raise
    
    def _format_txt_to_docx(self, input_path, output_path):
        """
        Converte um arquivo TXT para DOCX e aplica formatação ABNT.
        
        Args:
            input_path (str): Caminho do arquivo de entrada
            output_path (str): Caminho do arquivo de saída
            
        Returns:
            str: Caminho do arquivo formatado
        """
        try:
            # Ler conteúdo do arquivo TXT
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            # Criar documento DOCX
            doc = Document()
            
            # Configurar estilos ABNT
            self._apply_abnt_styles(doc)
            
            # Adicionar texto
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style = 'Normal'
            
            # Aplicar formatação ABNT ao conteúdo
            self._apply_abnt_formatting_to_docx(doc)
            
            # Salvar documento formatado
            doc.save(output_path)
            
            logger.info(f"Documento TXT convertido para DOCX com formatação ABNT: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Erro ao converter TXT para DOCX: {str(e)}")
            raise
    
    def _apply_abnt_styles(self, doc):
        """
        Aplica estilos ABNT ao documento.
        
        Args:
            doc: Documento DOCX
        """
        # Estilo Normal (corpo do texto)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.space_after = Pt(0)
        
        # Estilo de Título 1
        style = doc.styles['Heading 1']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.bold = True
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(12)
        
        # Estilo de Título 2
        style = doc.styles['Heading 2']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = True
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(6)
        
        # Estilo para citações longas
        if 'ABNT Citação' not in doc.styles:
            style = doc.styles.add_style('ABNT Citação', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(10)
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format.left_indent = Cm(4)
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
        
        # Estilo para referências
        if 'ABNT Referência' not in doc.styles:
            style = doc.styles.add_style('ABNT Referência', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(10)
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.left_indent = Cm(0)
            paragraph_format.first_line_indent = Cm(-0.75)
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(6)
    
    def _apply_abnt_formatting_to_docx(self, doc):
        """
        Aplica formatação ABNT ao conteúdo do documento DOCX.
        
        Args:
            doc: Documento DOCX
        """
        # Configurar margens do documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        
        # Processar parágrafos para identificar e formatar citações e referências
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Identificar citações longas (mais de 3 linhas)
            if len(text) > 240 and text.startswith('"') and text.endswith('"'):
                paragraph.style = 'ABNT Citação'
                
                # Atualizar formato de citação conforme ABNT 2023
                # Remover aspas das citações longas
                if text.startswith('"') and text.endswith('"'):
                    paragraph.text = text[1:-1]
                
                # Corrigir formato de autoria em citações
                paragraph.text = self._fix_citation_format(paragraph.text)
            
            # Identificar possíveis referências
            elif re.match(r'^[A-Z]+,\s+[A-Z]', text) or text.startswith('ASSOCIAÇÃO BRASILEIRA'):
                paragraph.style = 'ABNT Referência'
            
            # Processar citações em parágrafos normais
            else:
                # Corrigir formato de autoria em citações
                new_text = self._fix_citation_format(text)
                if new_text != text:
                    paragraph.text = new_text
    
    def _fix_citation_format(self, text):
        """
        Corrige o formato de citações conforme ABNT 2023.
        
        Args:
            text (str): Texto a ser corrigido
            
        Returns:
            str: Texto corrigido
        """
        # Padrão para citações entre parênteses com autoria em maiúsculas
        pattern = r'\(([A-Z]+)(,|\s+et\s+al\.)(.*?)\)'
        
        # Função para converter a primeira letra para maiúscula e o resto para minúscula
        def convert_case(match):
            author = match.group(1)
            separator = match.group(2)
            rest = match.group(3)
            
            # Converter autor para iniciar com maiúscula
            if author == author.upper():
                author = author.title()
            
            # Converter "ET AL." para "et al."
            if separator.lower() == ' et al.':
                separator = ' et al.'
            
            return f'({author}{separator}{rest})'
        
        # Aplicar correção
        corrected_text = re.sub(pattern, convert_case, text)
        
        # Corrigir expressões latinas para itálico (simulado com marcadores)
        latin_expressions = ['et al.', 'apud', 'in', 'loc. cit.', 'op. cit.', 'passim', 'sic']
        for expr in latin_expressions:
            # Não podemos aplicar itálico diretamente no texto, mas podemos marcar para processamento posterior
            corrected_text = corrected_text.replace(f' {expr} ', f' <i>{expr}</i> ')
            corrected_text = corrected_text.replace(f' {expr},', f' <i>{expr}</i>,')
            corrected_text = corrected_text.replace(f' {expr}.', f' <i>{expr}</i>.')
            corrected_text = corrected_text.replace(f' {expr})', f' <i>{expr}</i>)')
        
        return corrected_text

# Inicializar o formatador ABNT
abnt_formatter = ABNTFormatter()

@app.route('/')
def index():
    return render_template('index.html', logged_in='user_id' in session)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        if email in users_db and check_password_hash(users_db[email]['password'], password):
            session['user_id'] = users_db[email]['id']
            session['user_email'] = email
            flash('Login realizado com sucesso!')
            return redirect(url_for('index'))
        else:
            flash('Email ou senha incorretos.')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        name = request.form.get('name')
        
        if email in users_db:
            flash('Este email já está cadastrado.')
        else:
            user_id = str(uuid.uuid4())
            users_db[email] = {
                'id': user_id,
                'name': name,
                'password': generate_password_hash(password),
                'documents': []
            }
            
            session['user_id'] = user_id
            session['user_email'] = email
            flash('Cadastro realizado com sucesso!')
            return redirect(url_for('index'))
    
    return render_template('register.html')

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('user_email', None)
    flash('Você saiu da sua conta.')
    return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('Nenhum arquivo enviado')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('Nenhum arquivo selecionado')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        # Gerar nome de arquivo seguro e único
        original_filename = secure_filename(file.filename)
        filename = f"{uuid.uuid4()}_{original_filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        try:
            # Salvar o arquivo
            file.save(filepath)
            
            # Verificar se é um arquivo .gdoc
            if filepath.lower().endswith('.gdoc'):
                # Em uma versão real, usaríamos a API do Google Drive
                # Aqui, apenas simulamos a extração
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    try:
                        gdoc_data = json.load(f)
                        doc_id = gdoc_data.get('doc_id', 'unknown')
                    except:
                        doc_id = 'unknown'
                
                txt_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename}_extracted.txt")
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(f"Conteúdo extraído do Google Doc ID: {doc_id}\n\n")
                    f.write("Este é um documento simulado extraído de um arquivo Google Docs.\n")
                    f.write("Em um ambiente de produção, usaríamos a API do Google Drive para baixar o conteúdo real.")
                
                filepath = txt_path
                
            # Aplicar formatação ABNT
            output_path = abnt_formatter.format_document(filepath)
            
            # Gerar URL para download
            download_url = url_for('download_file', filename=os.path.basename(output_path))
            
            # Se o usuário estiver logado, salvar o documento no histórico
            if 'user_id' in session:
                user_email = session['user_email']
                if 'documents' not in users_db[user_email]:
                    users_db[user_email]['documents'] = []
                
                users_db[user_email]['documents'].append({
                    'original_name': original_filename,
                    'formatted_name': os.path.basename(output_path),
                    'date': 'Agora'  # Em produção, usar datetime
                })
            
            # Retornar para a página com link de download
            return render_template('index.html', 
                                  file_uploaded=True, 
                                  filename=original_filename, 
                                  download_url=download_url,
                                  logged_in='user_id' in session)
            
        except Exception as e:
            logger.error(f"Erro ao processar arquivo: {str(e)}")
            flash(f'Erro ao processar o arquivo: {str(e)}')
            return redirect(url_for('index'))
    else:
        flash('Tipo de arquivo não permitido')
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)
    except FileNotFoundError:
        flash('Arquivo não encontrado.')
        return redirect(url_for('index'))

@app.route('/send_email', methods=['POST'])
def send_email():
    if 'user_id' not in session:
        flash('Você precisa estar logado para enviar por email.')
        return redirect(url_for('login'))
    
    recipient_email = request.form.get('email')
    filename = request.form.get('filename')
    
    if not recipient_email or not filename:
        flash('Email ou arquivo não especificado.')
        return redirect(url_for('index'))
    
    try:
        # Simulação de envio de email (em produção, usar serviço real)
        flash(f'Email simulado enviado para {recipient_email} com o arquivo {filename}.')
        return redirect(url_for('index'))
        
    except Exception as e:
        logger.error(f"Erro ao enviar email: {str(e)}")
        flash(f'Erro ao enviar email: {str(e)}')
        return redirect(url_for('index'))

@app.route('/history')
def history():
    if 'user_id' not in session:
        flash('Você precisa estar logado para ver seu histórico.')
        return redirect(url_for('login'))
    
    user_email = session['user_email']
    documents = users_db[user_email].get('documents', [])
    
    return render_template('history.html', documents=documents)

if __name__ == '__main__':
    # Listening on 0.0.0.0 makes it accessible externally
    app.run(host='0.0.0.0', port=5000, debug=True)
